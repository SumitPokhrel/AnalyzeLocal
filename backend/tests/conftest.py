"""Shared test fixtures.

Test documents are generated at run time rather than committed as binary
files, so a contributor can read exactly what is in each one and why.
"""

from pathlib import Path

import docx
import pytest
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


@pytest.fixture(scope="session")
def documents(tmp_path_factory: pytest.TempPathFactory) -> dict[str, Path]:
    """Build one of each document the extractor has to handle."""
    directory = tmp_path_factory.mktemp("documents")
    paths: dict[str, Path] = {}
    paths.update(write_text_documents(directory))
    paths.update(write_pdf_documents(directory))
    paths.update(write_docx_documents(directory))
    return paths


def write_text_documents(directory: Path) -> dict[str, Path]:
    """Write plain text files in the encodings the extractor has to survive."""
    paths: dict[str, Path] = {}

    paths["utf8"] = directory / "utf8.txt"
    paths["utf8"].write_text(
        "Offer Letter\n\nBase salary: 145000 USD\n", encoding="utf-8"
    )

    paths["utf16"] = directory / "utf16.txt"
    paths["utf16"].write_text("Salary section\nBase: 145000\n", encoding="utf-16")

    # Raw cp1252 bytes. 0x93 and 0x94 are curly quotes there and are not
    # valid UTF-8. The length is even on purpose, because a utf-16 attempt
    # decodes any even length file into silent garbage instead of failing.
    paths["cp1252"] = directory / "cp1252.txt"
    paths["cp1252"].write_bytes(b"Bonus \x93target\x94 is 15%\n")

    # A UTF-8 byte order mark, which Windows editors add and which the plain
    # utf-8 codec leaves in the text as a stray character.
    paths["utf8_bom"] = directory / "utf8_bom.txt"
    paths["utf8_bom"].write_bytes(b"\xef\xbb\xbfSalary: 145000\n")

    paths["crlf"] = directory / "crlf.txt"
    paths["crlf"].write_text(
        "Line one   \r\nLine two\r\n\r\n\r\n\r\nLine three\r\n",
        encoding="utf-8",
        newline="",
    )

    paths["empty_text"] = directory / "empty.txt"
    paths["empty_text"].write_text("   \n\n  \n", encoding="utf-8")

    return paths


def write_pdf_documents(directory: Path) -> dict[str, Path]:
    """Write PDFs with a text layer, without one, and with encryption."""
    paths: dict[str, Path] = {}

    paths["pdf"] = directory / "text.pdf"
    page = canvas.Canvas(str(paths["pdf"]), pagesize=LETTER)
    page.drawString(72, 720, "Page one: base salary 145000 USD")
    page.drawString(72, 700, "Reports to Jane Doe")
    page.showPage()
    page.drawString(72, 720, "Page two: equity 4000 RSUs")
    page.showPage()
    page.save()

    # Pages with no text layer, standing in for a scan.
    paths["scanned_pdf"] = directory / "scanned.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with open(paths["scanned_pdf"], "wb") as handle:
        writer.write(handle)

    # A real password, which cannot be recovered.
    paths["locked_pdf"] = directory / "locked.pdf"
    with open(paths["locked_pdf"], "wb") as handle:
        encrypt_pdf(paths["pdf"], "secret").write(handle)

    # An empty user password, which is how many payroll exports arrive and
    # which should open without asking the user for anything.
    paths["empty_password_pdf"] = directory / "empty_password.pdf"
    with open(paths["empty_password_pdf"], "wb") as handle:
        encrypt_pdf(paths["pdf"], "", owner_password="owner").write(handle)

    return paths


def encrypt_pdf(source: Path, password: str, owner_password: str | None = None) -> PdfWriter:
    """Copy a PDF into a writer and encrypt it with the given password."""
    writer = PdfWriter()
    for page in PdfReader(str(source)).pages:
        writer.add_page(page)
    writer.encrypt(password, owner_password=owner_password)
    return writer


def write_docx_documents(directory: Path) -> dict[str, Path]:
    """Write Word documents, including one with a table between paragraphs."""
    paths: dict[str, Path] = {}

    # An offer letter shape: the numbers sit in a table in the middle of the
    # document, which is what makes document order worth testing.
    paths["docx"] = directory / "offer.docx"
    document = docx.Document()
    document.add_paragraph("Dear candidate,")
    document.add_paragraph("We are pleased to offer you the following terms.")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Component"
    table.cell(0, 1).text = "Amount"
    table.cell(1, 0).text = "Base salary"
    table.cell(1, 1).text = "145000 USD"
    table.cell(2, 0).text = "Signing bonus"
    table.cell(2, 1).text = "20000 USD"
    document.add_paragraph("Please sign and return by 2026-08-15.")
    document.save(str(paths["docx"]))

    paths["empty_docx"] = directory / "empty.docx"
    docx.Document().save(str(paths["empty_docx"]))

    # A docx extension on something that is not a zip archive at all.
    paths["fake_docx"] = directory / "fake.docx"
    paths["fake_docx"].write_bytes(b"this is not a zip archive")

    return paths
